"""Input adapters: turn document files into the stable JSON payloads nightward gates.

The gate core only ever sees JSON, so supporting a format means converting it
to a *stable* payload — that's all an adapter does. Two gating levels:

- **artifact** (`from_file`): sha256 + size of the raw bytes. Works for every
  format on earth (hwp, binaries, model weights) with zero dependencies, but
  re-saves that only churn metadata will breach it.
- **content** (`from_pdf` / `from_docx` / `from_xlsx` / `from_text`): structural
  facts plus a hash of the extracted content. Robust to byte-level noise,
  sensitive to what actually matters. Capture both when you want "the file
  changed" and "the content changed" reported separately.

Payloads stay small by design — hashes and counts, never raw document dumps —
so approved baselines remain reviewable in git.

`from_file`/`from_text` are stdlib-only. The rest lazy-import their parser and
raise a clear NightwardError pointing at `pip install "nightward[docs]"`.

Validated against real-world files (Korean PDF/XLSX/DOCX/HWP/cp949 TXT):
docs/experiments/2026-06-10-document-input-adapters.md.
"""
from __future__ import annotations

import hashlib
from pathlib import Path

from .errors import NightwardError

_DOCS_HINT = "pip install 'nightward[docs]'"


def _sha(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def _read(path: str | Path) -> bytes:
    p = Path(path)
    if not p.is_file():
        raise NightwardError(f"adapter input not found: {p}")
    return p.read_bytes()


def from_file(path: str | Path) -> dict:
    """Artifact fingerprint for any format at all (stdlib-only)."""
    data = _read(path)
    return {"sha256": _sha(data), "size_bytes": len(data)}


def from_text(path: str | Path, encodings: tuple[str, ...] = ("utf-8", "cp949")) -> dict:
    """Plain text with unknown encoding; falls back to the artifact fingerprint.

    Tries `encodings` in order (default covers UTF-8 and Korean legacy cp949).
    The text hash is over the *decoded* text re-encoded as UTF-8, so the same
    content stored in different encodings gates as equal.
    """
    raw = _read(path)
    for enc in encodings:
        try:
            text = raw.decode(enc)
        except UnicodeDecodeError:
            continue
        return {"encoding": enc, "chars": len(text), "lines": text.count("\n") + 1,
                "text_sha256": _sha(text.encode("utf-8"))}
    return {"encoding": "unknown"} | from_file(path)


def from_pdf(path: str | Path) -> dict:
    """Content-level PDF facts: page count + extracted-text hash.

    Robust to metadata churn (re-saves, /Producer stamps); breaches when the
    text content moves. Pair with `from_file` to also see byte-level changes.
    """
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise NightwardError(f"from_pdf needs pypdf - {_DOCS_HINT}") from exc
    _read(path)  # uniform not-found error before the parser's own
    reader = PdfReader(str(path))
    text = "\n".join((page.extract_text() or "") for page in reader.pages)
    return {"pages": len(reader.pages), "text_chars": len(text),
            "text_sha256": _sha(text.encode("utf-8"))}


def from_docx(path: str | Path) -> dict:
    """Content-level DOCX facts: paragraph/table counts + text hash."""
    try:
        import docx
    except ImportError as exc:
        raise NightwardError(f"from_docx needs python-docx - {_DOCS_HINT}") from exc
    _read(path)
    document = docx.Document(str(path))
    paragraphs = [p.text for p in document.paragraphs]
    text = "\n".join(paragraphs)
    return {"paragraphs": len(paragraphs), "tables": len(document.tables),
            "text_chars": len(text), "text_sha256": _sha(text.encode("utf-8"))}


def from_xlsx(path: str | Path) -> dict:
    """Content-level XLSX facts: per-sheet dimensions + a hash over cell values.

    Hashes computed values (data_only), not formulas or styling — a re-save
    that only touches zip timestamps or formatting gates as equal.
    """
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise NightwardError(f"from_xlsx needs openpyxl - {_DOCS_HINT}") from exc
    _read(path)
    workbook = load_workbook(str(path), read_only=True, data_only=True)
    sheets = {}
    h = hashlib.sha256()
    for ws in workbook.worksheets:
        rows = 0
        for row in ws.iter_rows(values_only=True):
            rows += 1
            h.update(repr(row).encode("utf-8"))
        sheets[ws.title] = {"rows": rows, "cols": ws.max_column}
    return {"sheets": sheets, "content_sha256": h.hexdigest()}
